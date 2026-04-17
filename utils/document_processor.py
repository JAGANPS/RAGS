"""
Document processing utilities for multiple file formats
"""
import os
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass
import hashlib

# Document loaders
from pypdf import PdfReader
from docx import Document as DocxDocument
import pandas as pd
import openpyxl


@dataclass
class ProcessedDocument:
    """Represents a processed document"""
    id: str
    filename: str
    content: str
    metadata: Dict[str, Any]
    file_type: str
    category: str  # SAM, ITAM, ITV, SRE
    chunks: List[Dict[str, Any]] = None

    def __post_init__(self):
        if self.chunks is None:
            self.chunks = []


class DocumentProcessor:
    """Process various document formats"""

    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.txt', '.csv', '.xlsx', '.xls'}

    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap

    def process_file(self, file_path: str, category: str = "general") -> ProcessedDocument:
        """Process a single file and return structured document"""
        path = Path(file_path)
        extension = path.suffix.lower()

        if extension not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file type: {extension}")

        content = self._extract_content(file_path, extension)
        doc_id = self._generate_id(file_path, content)

        doc = ProcessedDocument(
            id=doc_id,
            filename=path.name,
            content=content,
            metadata={
                "file_path": str(path.absolute()),
                "file_size": path.stat().st_size,
                "extension": extension,
            },
            file_type=extension[1:],  # Remove the dot
            category=category
        )

        # Create chunks
        doc.chunks = self._create_chunks(doc)
        return doc

    def _extract_content(self, file_path: str, extension: str) -> str:
        """Extract text content based on file type"""
        extractors = {
            '.pdf': self._extract_pdf,
            '.docx': self._extract_docx,
            '.doc': self._extract_docx,  # Try docx parser
            '.txt': self._extract_txt,
            '.csv': self._extract_csv,
            '.xlsx': self._extract_excel,
            '.xls': self._extract_excel,
        }

        extractor = extractors.get(extension)
        if extractor:
            return extractor(file_path)
        raise ValueError(f"No extractor for {extension}")

    def _extract_pdf(self, file_path: str) -> str:
        """Extract text from PDF"""
        reader = PdfReader(file_path)
        text_parts = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_parts.append(f"[Page {page_num + 1}]\n{text}")
        return "\n\n".join(text_parts)

    def _extract_docx(self, file_path: str) -> str:
        """Extract text from DOCX"""
        doc = DocxDocument(file_path)
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        return "\n\n".join(paragraphs)

    def _extract_txt(self, file_path: str) -> str:
        """Extract text from TXT"""
        with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
            return f.read()

    def _extract_csv(self, file_path: str) -> str:
        """Extract text from CSV"""
        df = pd.read_csv(file_path)
        # Convert to readable format
        lines = [f"Columns: {', '.join(df.columns)}"]
        for idx, row in df.iterrows():
            row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
            lines.append(f"Row {idx + 1}: {row_text}")
        return "\n".join(lines)

    def _extract_excel(self, file_path: str) -> str:
        """Extract text from Excel"""
        xls = pd.ExcelFile(file_path)
        all_text = []
        for sheet_name in xls.sheet_names:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            all_text.append(f"=== Sheet: {sheet_name} ===")
            all_text.append(f"Columns: {', '.join(map(str, df.columns))}")
            for idx, row in df.iterrows():
                row_text = " | ".join([f"{col}: {val}" for col, val in row.items()])
                all_text.append(f"Row {idx + 1}: {row_text}")
        return "\n".join(all_text)

    def _generate_id(self, file_path: str, content: str) -> str:
        """Generate unique document ID"""
        hash_input = f"{file_path}:{content[:500]}"
        return hashlib.sha256(hash_input.encode()).hexdigest()[:16]

    def _create_chunks(self, doc: ProcessedDocument) -> List[Dict[str, Any]]:
        """Create overlapping chunks from document content"""
        content = doc.content
        chunks = []

        if len(content) <= self.chunk_size:
            chunks.append({
                "chunk_id": f"{doc.id}_0",
                "text": content,
                "start_idx": 0,
                "end_idx": len(content),
                "metadata": {
                    "doc_id": doc.id,
                    "filename": doc.filename,
                    "category": doc.category,
                    "chunk_index": 0
                }
            })
        else:
            start = 0
            chunk_index = 0
            while start < len(content):
                end = min(start + self.chunk_size, len(content))

                # Try to break at sentence boundary
                if end < len(content):
                    for sep in ['. ', '.\n', '\n\n', '\n', ' ']:
                        last_sep = content[start:end].rfind(sep)
                        if last_sep > self.chunk_size // 2:
                            end = start + last_sep + len(sep)
                            break

                chunk_text = content[start:end].strip()
                if chunk_text:
                    chunks.append({
                        "chunk_id": f"{doc.id}_{chunk_index}",
                        "text": chunk_text,
                        "start_idx": start,
                        "end_idx": end,
                        "metadata": {
                            "doc_id": doc.id,
                            "filename": doc.filename,
                            "category": doc.category,
                            "chunk_index": chunk_index
                        }
                    })
                    chunk_index += 1

                start = end - self.chunk_overlap
                if start >= len(content) - self.chunk_overlap:
                    break

        return chunks

    def process_directory(self, dir_path: str, category: str = "general") -> List[ProcessedDocument]:
        """Process all supported files in a directory"""
        documents = []
        path = Path(dir_path)

        if not path.exists():
            return documents

        for file_path in path.rglob("*"):
            if file_path.suffix.lower() in self.SUPPORTED_EXTENSIONS:
                try:
                    doc = self.process_file(str(file_path), category)
                    documents.append(doc)
                except Exception as e:
                    print(f"Error processing {file_path}: {e}")

        return documents
